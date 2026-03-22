"""Generate PropVal Database Inventory Word document."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading("PropVal Complete Database Inventory", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Generated: 22 March 2026", style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ── Overview ──
doc.add_heading("Overview: 3 Database Tiers", level=1)
t = doc.add_table(rows=4, cols=5, style="Light Grid Accent 1")
for i, row in enumerate([
    ["Tier", "Technology", "Location", "Size", "Purpose"],
    ["Local", "DuckDB", "Research/EPC PPD merge project/db/propval.duckdb", "1.9 GB", "PPD+EPC merged London data (1.1M matched)"],
    ["Local", "SQLite x3", "backend/data/", "3.5 GB", "UPRN coords (41.5M), INSPIRE polygons (2M), Leases (4.3M)"],
    ["Cloud", "PostgreSQL", "Supabase", "\u2014", "Cases, cache tables, firm/user data, reports (27 tables)"],
]):
    for j, val in enumerate(row):
        t.rows[i].cells[j].text = val
        if i == 0:
            for run in t.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True

# ── A. DuckDB ──
doc.add_heading("A. LOCAL: DuckDB (propval.duckdb)", level=1)
doc.add_heading("Tables", level=2)
t = doc.add_table(rows=5, cols=3, style="Light Grid Accent 1")
for i, row in enumerate([
    ["Table", "Rows", "Purpose"],
    ["matched", "1,136,116", "PPD+EPC successfully linked transactions"],
    ["unmatched", "101,721", "PPD records without EPC match"],
    ["epc", "4,436,405", "Raw EPC certificates (all London)"],
    ["construction_age", "4,436,405", "Derived building age estimates"],
]):
    for j, val in enumerate(row):
        t.rows[i].cells[j].text = val
        if i == 0:
            for run in t.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True

doc.add_heading("matched table (key columns)", level=2)
cols_data = [
    ["Column", "Type", "Source", "Notes"],
    ["transaction_id", "TEXT", "PPD", "Unique sale identifier"],
    ["price", "INTEGER", "PPD", "Sale price"],
    ["date_of_transfer", "DATE", "PPD", "Sale date"],
    ["postcode", "TEXT", "PPD", ""],
    ["outward_code", "TEXT", "PPD", "First half of postcode"],
    ["saon", "TEXT", "PPD", "Flat number (e.g. FLAT 25)"],
    ["paon", "TEXT", "PPD", "House number/building name"],
    ["street", "TEXT", "PPD", ""],
    ["ppd_type", "TEXT", "PPD", "D/S/T/F/O"],
    ["duration", "TEXT", "PPD", "F=freehold, L=leasehold"],
    ["old_new", "TEXT", "PPD", "Y=new build"],
    ["ppd_category", "TEXT", "PPD", "A=standard, B=non-standard"],
    ["UPRN", "TEXT", "EPC", "WARNING: Can be building-level"],
    ["LMK_KEY", "TEXT", "EPC", "EPC certificate key"],
    ["ADDRESS1", "TEXT", "EPC", "Free-text address line"],
    ["epc_type", "TEXT", "EPC", "Property type from EPC"],
    ["TOTAL_FLOOR_AREA", "REAL", "EPC", "Floor area sqm"],
    ["CURRENT_ENERGY_RATING", "TEXT", "EPC", "A-G"],
    ["CONSTRUCTION_AGE_BAND", "TEXT", "EPC", ""],
    ["lat", "REAL", "OS Open UPRN", "WGS84"],
    ["lon", "REAL", "OS Open UPRN", "WGS84"],
    ["match_tier", "TEXT", "System", "Which matching tier succeeded"],
    ["match_confidence", "REAL", "System", "0.0-1.0"],
]
t = doc.add_table(rows=len(cols_data), cols=4, style="Light Grid Accent 1")
for i, row in enumerate(cols_data):
    for j, val in enumerate(row):
        t.rows[i].cells[j].text = val
        if i == 0:
            for run in t.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True

doc.add_heading("The UPRN Problem", level=2)
doc.add_paragraph("~11,430 transactions (1%) have building-level UPRNs shared across multiple flats.")
doc.add_paragraph("Worst case: 1 UPRN maps to 28 different flats.")
doc.add_paragraph("Root cause: EPC/OS assigns one UPRN per building, merge pipeline inherits it.")
p = doc.add_paragraph("Needs fundamental fix at data layer, not patching each query.")
p.runs[0].bold = True

# ── B. SQLite ──
doc.add_heading("B. LOCAL: SQLite Databases", level=1)
for name, size, table, rows, pk, cols in [
    ("uprn_coords.db", "1.3 GB", "uprn_coords", "41,466,127", "uprn (INTEGER)", "uprn, lat (REAL), lon (REAL) \u2014 OS Open UPRN, ~1-5m accuracy"),
    ("inspire_polygons.db", "1.8 GB", "polygons", "2,061,222", "inspire_id (INTEGER)", "inspire_id, geojson (TEXT) \u2014 INSPIRE property boundaries"),
    ("leases.db", "391 MB", "registered_leases", "4,280,569", "uprn (TEXT)", "uprn, date_of_lease, term_years, expiry_date"),
]:
    doc.add_heading(f"{name} ({size})", level=2)
    t = doc.add_table(rows=2, cols=3, style="Light Grid Accent 1")
    for j, val in enumerate(["Table", "Rows", "PK"]):
        t.rows[0].cells[j].text = val
        for run in t.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
    for j, val in enumerate([table, rows, pk]):
        t.rows[1].cells[j].text = val
    doc.add_paragraph(f"Columns: {cols}")

doc.add_heading("inspire_centroids_london.json (121 MB)", level=2)
doc.add_paragraph("2.06M INSPIRE centroids loaded to memory at startup. Used for KDTree nearest-centroid lookups.")

# ── C. Supabase ──
doc.add_heading("C. CLOUD: Supabase (27 tables)", level=1)

tiers = [
    ("Tier 1: Property Library (shared, read-only)", [
        ["properties", "uprn TEXT", "\u2014", "UPRN master record"],
        ["property_enrichment", "UUID", "uprn \u2192 properties", "Cached API responses per source"],
        ["property_enrichment_history", "UUID", "\u2014", "Archive of enrichment changes"],
        ["comparable_transactions", "UUID", "\u2014", "Comparable transaction cache"],
        ["outward_code_adjacency", "(outward, adjacent)", "\u2014", "Adjacent postcode lookup"],
    ]),
    ("Tier 2: Cache (backend service-role only)", [
        ["price_paid_cache", "transaction_id TEXT", "\u2014", "Bulk PPD data per outward code"],
        ["ppd_cache_status", "outward_code TEXT", "\u2014", "PPD cache freshness tracker"],
        ["epc_cache", "lmk_key TEXT", "\u2014", "Bulk EPC data per outward code"],
        ["epc_cache_status", "outward_code TEXT", "\u2014", "EPC cache freshness tracker"],
    ]),
    ("Tier 3: Cases (firm-private, RLS)", [
        ["cases", "UUID", "surveyor_id, firm_id \u2192 firms", "Surveyor working files"],
        ["property_snapshots", "UUID", "based_on_id \u2192 self, firm_id \u2192 firms", "Immutable property data snapshots"],
        ["case_comps", "UUID", "case_id \u2192 cases, snapshot_id \u2192 snapshots", "Adopted comparables junction"],
        ["case_comparables", "UUID", "transaction_id \u2192 comparable_transactions", "Legacy comps (deprecated)"],
    ]),
    ("Tier 4: Firm & Organization", [
        ["firms", "UUID", "\u2014", "Multi-tenancy firm records"],
        ["firm_members", "UUID", "firm_id \u2192 firms", "Membership junction"],
        ["firm_templates", "UUID", "firm_id \u2192 firms", "Report boilerplate + AI prompts"],
        ["firm_signatories", "UUID", "\u2014", "Staff details for report signing"],
    ]),
    ("Tier 5: Reports & Review", [
        ["report_templates", "UUID", "created_by \u2192 auth.users", "ARTG template definitions"],
        ["report_copies", "UUID", "case_id \u2192 cases", "Immutable versioned report snapshots"],
        ["qa_results", "UUID", "copy_id \u2192 report_copies", "AI QA findings"],
        ["review_requests", "UUID", "copy_id \u2192 report_copies", "Review assignments"],
        ["review_events", "UUID", "review_id \u2192 review_requests", "Audit log"],
        ["notifications", "UUID", "user_id \u2192 auth.users", "In-app notifications"],
    ]),
    ("Tier 6: Market Intelligence", [
        ["news_articles", "UUID", "\u2014", "RSS feed articles"],
        ["macro_indicators", "UUID", "\u2014", "Base rate, CPI, HPI, etc."],
    ]),
]

for tier_name, rows_data in tiers:
    doc.add_heading(tier_name, level=2)
    t = doc.add_table(rows=len(rows_data) + 1, cols=4, style="Light Grid Accent 1")
    for j, val in enumerate(["Table", "PK", "Key FKs", "Purpose"]):
        t.rows[0].cells[j].text = val
        for run in t.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            t.rows[i + 1].cells[j].text = val

doc.add_heading("Supabase RPC Functions", level=2)
t = doc.add_table(rows=3, cols=2, style="Light Grid Accent 1")
for j, val in enumerate(["Function", "Purpose"]):
    t.rows[0].cells[j].text = val
    for run in t.rows[0].cells[j].paragraphs[0].runs:
        run.bold = True
t.rows[1].cells[0].text = "get_user_firm_id()"
t.rows[1].cells[1].text = "Returns user firm_id (prevents RLS recursion)"
t.rows[2].cells[0].text = "autocomplete_by_postcode(pc)"
t.rows[2].cells[1].text = "Fast EPC address lookup with 5s timeout"

# ── D. Data Flow ──
doc.add_heading("D. Data Flow: How UPRN Connects Everything", level=1)
flow_lines = [
    "User searches postcode",
    "  \u2192 EPC API / epc_cache / DuckDB autocomplete \u2192 address + UPRN",
    "  \u2192 UPRN \u2192 uprn_coords.db \u2192 lat/lon",
    "  \u2192 UPRN \u2192 inspire_polygons.db \u2192 boundary polygon",
    "  \u2192 UPRN \u2192 DuckDB matched \u2192 sale history + EPC data",
    "  \u2192 UPRN \u2192 price_paid_cache \u2192 additional sale history",
    "  \u2192 UPRN \u2192 leases.db \u2192 lease details",
    "  \u2192 UPRN \u2192 property_enrichment \u2192 cached API responses",
    "  \u2192 lat/lon \u2192 Environment Agency, planning.data.gov, etc.",
]
for line in flow_lines:
    p = doc.add_paragraph(line)
    for run in p.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)

doc.add_paragraph()
p = doc.add_paragraph("UPRN is the universal spine. If UPRN is wrong (building-level), every downstream lookup is contaminated.")
p.runs[0].bold = True
p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.save("docs/PropVal_Database_Inventory.docx")
print("Saved to docs/PropVal_Database_Inventory.docx")
