from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# --- Helper functions ---
def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()  # spacing after table

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Setting Up a Residential\nValuation Firm in London', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Complete Setup Guide')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.add_paragraph()
context = doc.add_paragraph()
context.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = context.add_run('Prepared for PropVal Pilot Strategy\nMarch 2025')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# CONTEXT
# ============================================================
add_heading('Context', 1)
add_para('Start own RICS-regulated valuation firm as PropVal\'s pilot vehicle. 3-person team (2 MRICS valuers rotating inspections/writing, 1 admin). Target 600 reports/year at £700 avg = £420k gross. Firm operates as separate Ltd company alongside AllRange/PropVal Tech Ltd.')

doc.add_page_break()

# ============================================================
# PHASE 1
# ============================================================
add_heading('Phase 1: Company Formation (Week 1-2)', 1)

add_heading('Register Ltd Company', 2)
add_bullet('Companies House: £12 online, Certificate of Incorporation within 24 hours', 'Companies House: ')
add_bullet('SIC codes: 68310 (Real estate agencies) or 71129 (Other engineering activities — surveying)')
add_bullet('From November 2025: mandatory identity verification for all directors')
add_bullet('You as director, wife as director/company secretary')
add_bullet('Registered office can be home address initially')

add_heading('HMRC Registration', 2)
add_bullet('Register within 3 months of starting trading', 'Corporation Tax: ')
add_bullet('Small profits rate: 19% (up to £50k), main rate 25% (over £250k)')
add_bullet('Register before first payday (takes up to 15 working days)', 'PAYE: ')
add_bullet('Employer NIC: 15% from April 2025')
add_bullet('£10,500/yr off employer\'s NI (NOT available if sole employee is a director)', 'Employment Allowance: ')

add_heading('Business Bank Account', 2)
add_bullet('Required for Ltd company')
add_bullet('Starling, Monzo (free, fast digital setup) or traditional bank')
add_bullet('Need: Certificate of Incorporation, photo ID, proof of address')

add_heading('VAT Considerations', 2)
add_bullet('Threshold: £90,000 (may change April 2026)')
add_bullet('Lender/panel valuations are VAT exempt (treated as part of financial services)', 'CRITICAL: ')
add_bullet('Private valuations are standard-rated (20%)')
add_bullet('Mixed supplies create partial exemption complexity — get specialist accountant advice')

doc.add_page_break()

# ============================================================
# PHASE 2
# ============================================================
add_heading('Phase 2: RICS & Professional Registration (Week 2-4)', 1)

add_heading('RICS Firm Registration', 2)
add_bullet('Must register as RICS Regulated Firm to trade as chartered surveying practice')
add_bullet('Requirement: at least one MRICS/FRICS principal')
add_bullet('Apply via My Account portal — 5 steps (eligibility, registrations, application with PII & CHP details, declaration, payment)')
add_bullet('Email: regulation@rics.org')
add_bullet('MRICS personal subscription ~£578/yr (2026), Firm regulation ~£100-225/yr', 'Fees: ')

add_heading('Valuer Registration Scheme (VRS)', 2)
add_bullet('Mandatory for all RICS members doing valuations — individual registration, not firm')
add_bullet('If you achieved Valuation competency Level 3 at APC → register directly via My Account, no additional assessment')
add_bullet('If not: Valuer Registration Assessment required (application, up to 100 days supervised experience, case study, CPD record)')
add_bullet('Annual renewal, periodic RICS audit checking Red Book compliance')
add_bullet('Fee: ~£100-140/yr')
add_bullet('The other MRICS must also be VRS registered')

add_heading('Red Book Compliance (effective 31 Jan 2025)', 2)
add_bullet('All valuations must comply with RICS Valuation — Global Standards')
add_bullet('Mandatory: written terms of engagement (VPS 1), report content standards (VPS 6)')
add_bullet('Thorough inspection, comparable evidence')
add_bullet('ESG now mandatory', 'NEW: ')
add_bullet('AI/AVM outputs only count as valuation if professional judgement applied')
add_bullet('Must declare and manage conflicts of interest')

add_heading('CPD', 2)
add_bullet('20 hours/year per RICS member (minimum 10 formal/structured)')
add_bullet('Must log and evidence — RICS audits compliance')
add_bullet('Relevant: Red Book updates, AML, market updates, technical surveying, ethics')

doc.add_page_break()

# ============================================================
# PHASE 3
# ============================================================
add_heading('Phase 3: Insurance (Week 3-4)', 1)

add_heading('Professional Indemnity Insurance (PII) — MANDATORY', 2)
add_para('RICS minimums by fee income:', bold=True)
add_table(['Fee Income', 'Minimum Cover'], [
    ['Under £100k', '£250,000'],
    ['£100k-£200k', '£500,000'],
    ['Over £200k', '£1,000,000'],
])
add_bullet('Panel requirements: Most lenders require £1m minimum, some require £5m-£10m')
add_bullet('Must be "each and every claim" basis with full retroactive cover')
add_bullet('Must be from RICS Listed Insurer, complying with RICS Approved Minimum Wording')
add_bullet('~£2,000-5,000/yr for a new small firm', 'Cost: ')
add_bullet('6 years minimum if firm ceases practice (consumer run-off automatic from July 2025)', 'Run-off cover: ')
add_bullet('Howden, Hiscox, PIB Insurance, PI Expert, Simcox Brokers, Lockton, PolicyBee', 'Providers: ')

add_heading('Employer\'s Liability Insurance — LEGALLY REQUIRED', 2)
add_bullet('Minimum cover: £5 million (most policies are £10m)')
add_bullet('Must display certificate (or make electronically accessible)')
add_bullet('Fine: up to £2,500 per day for not having it')
add_bullet('Cost: ~£400-800/yr')

add_heading('Public Liability Insurance', 2)
add_bullet('Not legally required but effectively essential')
add_bullet('Typical cover: £1m-£5m, Cost: ~£200-500/yr')
add_bullet('Many panels require it')

add_heading('Other Insurance', 2)
add_bullet('~£300-800/yr (increasingly expected)', 'Cyber Insurance: ')
add_bullet('Business use cover essential for inspection vehicles', 'Motor Insurance: ')
add_bullet('Business use endorsement if working from home (~£30-80/yr extra)', 'Home Insurance: ')

doc.add_page_break()

# ============================================================
# PHASE 4
# ============================================================
add_heading('Phase 4: Compliance & Policies (Week 4-6)', 1)

add_heading('AML (Anti-Money Laundering)', 2)
add_bullet('RICS is your AML supervisory body (no separate HMRC registration needed)')
add_para('Required documentation:', bold=True)
add_bullet('Written firm-wide AML policy')
add_bullet('Firm-wide risk assessment (review annually)')
add_bullet('Customer Due Diligence (CDD) procedures — verify identity before every business relationship')
add_bullet('Appoint MLRO (Money Laundering Reporting Officer) — in small firm, this is you')
add_bullet('SAR (Suspicious Activity Report) procedures → file with NCA')
add_bullet('Record keeping: all CDD records retained 5 years after end of relationship')
add_bullet('Sanctions screening: HM Treasury consolidated list, UN sanctions')
add_bullet('Staff training: at start + annual refresher, record all training')
p = add_para('⚠ Tipping off is a criminal offence — never tell client about SAR filing', bold=True)

add_heading('Complaints Handling Procedure (CHP) — RICS MANDATORY', 2)
add_bullet('Named handler (not person complained about)')
add_bullet('Acknowledge within 7 days')
add_bullet('Substantive response within 28 days')
add_bullet('Two-stage internal process')
add_bullet('Info about independent redress scheme at each stage')
add_bullet('Must provide to clients at point of instruction (in terms of engagement)')
add_bullet('Maintain complaints log — RICS may inspect')
add_bullet('Independent redress: CEDR (Centre for Effective Dispute Resolution)')

add_heading('Data Protection', 2)
add_bullet('Mandatory, ~£40/yr (Tier 1, online at ico.org.uk)', 'ICO Registration: ')
add_bullet('Failure to register: criminal offence, fine up to £4,350')
add_para('Required documentation:', bold=True)
add_bullet('Privacy Notice (to clients, property occupiers, employees)')
add_bullet('Records of Processing Activities (ROPA)')
add_bullet('Data Processing Agreements with third-party processors')
add_bullet('Data breach procedure (report to ICO within 72 hours if risk)')
add_bullet('Subject Access Request procedure (respond within 1 calendar month)')

add_para('Retention periods:', bold=True)
add_table(['Record Type', 'Retention Period'], [
    ['Valuation reports/working files', '6yr minimum, 15yr recommended'],
    ['AML/CDD records', '5 years after end of relationship'],
    ['Employee records', '6 years after employment ends'],
    ['Financial/tax records', '6 years'],
])

add_heading('Health & Safety', 2)
add_para('Lone Worker Policy — critical for surveyors:', bold=True)
add_bullet('Check-in/check-out procedures for every property visit')
add_bullet('Escalation protocol for failed check-ins')
add_bullet('Lone worker app (StaySafe, PeopleSafe — ~£5-15/user/month)')
add_bullet('Right to withdraw from unsafe situations')
add_bullet('Generic risk assessment for property inspections')
add_bullet('Dynamic risk assessment training')

add_heading('Employment Setup', 2)
add_bullet('Written statement of employment on or before day 1', 'Contracts: ')
add_bullet('Mandatory from day 1 — 8% minimum total (3% employer, 5% employee), NEST easiest provider', 'Pension auto-enrolment: ')
add_bullet('Must verify before employment starts — criminal offence to skip', 'Right to work checks: ')
add_bullet('Basic check (£18-26) — many panels require it', 'DBS checks: ')

doc.add_page_break()

# ============================================================
# PHASE 5
# ============================================================
add_heading('Phase 5: Technology & Equipment (Week 5-7)', 1)

add_heading('Comparable Evidence Subscriptions', 2)
add_table(['Service', 'Cost', 'Purpose'], [
    ['Rightmove Plus', '~£150-250/month', 'Current asking prices, listing history'],
    ['Land Registry', '£3-6 per search / free bulk', 'Title info, transaction data'],
    ['EPC Register', 'Free', 'Energy performance data'],
    ['OS Mapping (Promap)', '~£100-200/yr', 'Site plans'],
    ['BCIS', '~£200-400/yr', 'Reinstatement cost data'],
    ['Flood data', '~£20-40 per search', 'Flood risk reports'],
])

add_heading('Surveying Equipment', 2)
add_table(['Item', 'Cost'], [
    ['Laser measure (Leica DISTO / Bosch GLM)', '£80-250'],
    ['Moisture meter (Protimeter)', '£150-400'],
    ['Tablet/laptop for on-site', '£500-1,500'],
    ['Spirit level, steel tape (backup)', '£30-50'],
    ['Torch', '£20-50'],
    ['PPE kit (boots, hi-vis, hard hat, masks, gloves)', '£100-200'],
    ['First aid kit (vehicle)', '£20-30'],
])

add_heading('Software', 2)
add_table(['Service', 'Cost'], [
    ['Accounting (Xero/QuickBooks)', '£15-42/month'],
    ['Microsoft 365 Business', '£9.40/user/month'],
    ['Payroll (HMRC Basic PAYE Tools)', 'Free (up to 9 employees)'],
    ['PropVal', '£35/seat/month + £35/case'],
])

add_heading('Vehicle', 2)
add_bullet('Essential for London residential surveying (outer boroughs especially)')
add_bullet('Must be ULEZ compliant (Euro 6 petrol ~2015+, Euro 6 diesel ~2016+)')
add_bullet('ULEZ: £12.50/day if non-compliant (covers all Greater London)')
add_bullet('Congestion Charge: £15/day (Mon-Fri 7am-6pm, Sat-Sun 12pm-6pm, Central London)')
add_bullet('Running costs: ~£300-500/month')
add_bullet('Business mileage: 45p/mile first 10,000, 25p thereafter')

add_heading('Professional Presence', 2)
add_bullet('~£10-15/yr for .co.uk + business email (never use gmail/hotmail)', 'Domain + email: ')
add_bullet('Essential — services, qualifications, CHP, privacy notice. DIY £12-30/month or professional £500-3,000', 'Website: ')
add_bullet('Included in RICS membership', 'RICS Find a Surveyor: ')
add_bullet('Free, essential for local search', 'Google Business Profile: ')

doc.add_page_break()

# ============================================================
# PHASE 6
# ============================================================
add_heading('Phase 6: Panel Applications (Week 7-8, then 4-12 weeks)', 1)

add_heading('How the Market Works', 2)
add_para('Lender → Panel Manager/AMC → Your Firm → Individual Valuer', bold=True)
add_para('Almost NO major UK lenders manage their own panels. You apply to the AMC, not the lender.')

add_heading('Priority Application Order', 2)

add_heading('1. Method Valuation Management — APPLY FIRST', 3)
add_bullet('Access to 50+ lenders, hundreds of monthly instructions')
add_bullet('You keep 100% of your fee (no commission — major differentiator)', 'KEY: ')
add_bullet('Requirements: RICS Regulated firm, 2+ directors, 2+ MRICS/FRICS, £1m min PII, 2yr+ PQE')
add_bullet('5-stage algorithmic allocation, cloud-based platform (Method xi)')
add_bullet('Register: method-vm.co.uk/valuers/register or helpdesk@method-vm.co.uk / 01642 269 306')

add_heading('2. Countrywide Surveying Services (Connells) — HIGHEST VOLUME', 3)
add_bullet('700+ staff, 450+ in-house surveyors, 50+ affiliated firms')
add_bullet('Lenders: Nationwide, HSBC, Santander, Leeds BS, Skipton BS, 40+ total')
add_bullet('Apply: connells-surveyors.co.uk/survey-valuation-panel/')
add_bullet('Panel Manager admin fee: £26 per valuation')

add_heading('3. e.surv (LSL Property Services) — SECOND LARGEST', 3)
add_bullet('600+ RICS surveyors')
add_bullet('Lenders: Virgin Money (sole), Halifax/Lloyds, Nationwide, Yorkshire BS, Danske Bank')
add_bullet('Primarily employed surveyors — external panel for coverage gaps')
add_bullet('Contact: recruitment@esurv.co.uk')

add_heading('4. Gateway Surveyors', 3)
add_bullet('All high-street lenders, building societies, equity release')
add_bullet('1,500+ individual valuers, ISO 9001:2015')

add_heading('5. CVN Services', 3)
add_bullet('Acts purely for lenders, portal + API connections')

add_heading('6. SDL Surveying / SDL Network', 3)
add_bullet('Network for independents, guaranteed work allocation, free CPD')
add_bullet('Apply: careers.sdlsurveying.co.uk')

add_heading('7. VAS Valuation Group — REGISTER INTEREST (currently closed)', 3)
add_bullet('200+ panel firms, 150+ lenders')
add_bullet('Register interest: vasmgt@vas-group.co.uk')

add_heading('8. Specialist/Bridging Lenders — EASIEST ENTRY POINT', 3)
add_bullet('AWH panel: 70+ non-bank lenders (Close Brothers, West One, LendInvest, Shawbrook, OakNorth, Glenhawk, etc.)')
add_bullet('Lower barriers, often higher per-job fees')

add_heading('Lender → AMC Mapping', 2)
add_table(['Lender', 'Panel Manager(s)'], [
    ['Nationwide BS', 'Countrywide (lead) + e.surv'],
    ['HSBC', 'Countrywide'],
    ['Santander', 'Countrywide (lead, 5-year deal)'],
    ['Halifax/Lloyds/BOS', 'e.surv + L&G Surveying'],
    ['Barclays', 'L&G Surveying + Connells + e.surv'],
    ['NatWest', 'L&G Surveying + panel managers'],
    ['Virgin Money/Clydesdale', 'e.surv (sole)'],
    ['Yorkshire BS', 'e.surv'],
    ['Leeds BS', 'Countrywide'],
    ['Skipton BS', 'Connells S&V'],
])

add_heading('Application Requirements', 2)
add_bullet('MRICS or FRICS (most exclude AssocRICS)')
add_bullet('VRS registered')
add_bullet('RICS Regulated Firm status')
add_bullet('Minimum £1m PII (some require £2m-£5m)')
add_bullet('Minimum 2 directors/partners and 2 MRICS/FRICS surveyors')
add_bullet('Minimum 2yr PQE')
add_bullet('Quality is critical — this is where firms get rejected', 'Sample reports (2-3, anonymised): ')
add_bullet('DBS check + AML compliance evidence')
add_bullet('Geographic coverage specification (London boroughs/postcodes)')

add_heading('How Allocation Works', 2)
add_bullet('Geographic proximity/postcode (primary factor)')
add_bullet('Surveyor availability/capacity')
add_bullet('Property type match')
add_bullet('PI cover vs property value')
add_bullet('Performance history/quality scores (RAG rating)')
add_bullet('Fee competitiveness (some panels)')

add_para('Getting more work:', bold=True)
add_bullet('Faster turnaround, fewer PVQs')
add_bullet('Wider geographic coverage')
add_bullet('Specialist competencies (HMO, BTL, new build, leasehold, equity release)')

add_heading('SLAs', 2)
add_bullet('Inspection: typically within 48 hours of instruction')
add_bullet('Report: same day as inspection')
add_bullet('Desktop valuations: within 24-48 hours')

add_heading('Payment Terms', 2)
add_bullet('Panels typically pay 30-45 days after report submission')
add_bullet('Some pay monthly in arrears')
add_bullet('Need 2-3 months working capital before revenue flows', 'IMPORTANT: ')

doc.add_page_break()

# ============================================================
# PHASE 7
# ============================================================
add_heading('Phase 7: Office & Final Setup (Week 6-8)', 1)

add_heading('Working from Home — Viable Initially', 2)
add_bullet('Check mortgage/lease: may need consent for business use')
add_bullet('Home insurance: business use endorsement (~£30-80/yr)')
add_bullet('Planning: generally not needed if character of dwelling unchanged')
add_bullet('Claim proportion of household costs against tax (simplified: £6/week)')
add_bullet('Avoid dedicating a room exclusively to business (loses CGT principal residence relief)', '⚠ ')
add_bullet('Virtual office for professional registered address: £15-50/month')

doc.add_page_break()

# ============================================================
# FINANCIAL SUMMARY
# ============================================================
add_heading('Financial Summary', 1)

add_heading('Startup Costs', 2)
add_table(['Category', 'Cost'], [
    ['Company formation', '£12-50'],
    ['RICS firm registration + VRS (×2)', '£750-1,000'],
    ['RICS personal subscriptions (×2)', '~£1,156/yr'],
    ['PII', '£2,000-5,000'],
    ['Employer\'s Liability', '£400-800'],
    ['Public Liability', '£200-500'],
    ['ICO Registration', '£40'],
    ['Surveying equipment (×2 sets)', '£800-2,000'],
    ['Laptop/tablets (×2)', '£1,000-3,000'],
    ['Website', '£500-2,000'],
    ['DBS checks (×2)', '£36-52'],
    ['AML training', '£50-200'],
    ['TOTAL (without vehicle)', '£7,000-15,000'],
    ['Vehicle (ULEZ-compliant used)', '£5,000-15,000'],
])

add_heading('Monthly Operating Costs (3 people)', 2)
add_table(['Category', 'Monthly'], [
    ['Insurance (PII + EL + PL amortised)', '£220-530'],
    ['RICS subscriptions (amortised)', '£175-225'],
    ['Rightmove Plus + data subs', '£250-400'],
    ['Software (accounting, Office, PropVal)', '£150-250'],
    ['Vehicle running costs', '£300-500'],
    ['Phone/internet', '£100-150'],
    ['Accountant', '£100-200'],
    ['Payroll (admin salary + employer NI + pension)', '£2,500-3,500'],
    ['TOTAL', '~£3,800-5,750/month'],
])

add_heading('Cash Flow Projection', 2)
add_table(['Period', 'Revenue', 'Costs', 'Cumulative'], [
    ['Month 1-2 (setup)', '£0', '~£23k', '-£23k'],
    ['Month 3-4 (first instructions)', '£0-5k', '£8k', '-£26k to -£31k'],
    ['Month 5-6 (ramping)', '£15-22k', '£8k', 'Recovering'],
    ['Month 7-12 (steady state)', '£33k/month', '£5k/month', 'Profitable'],
    ['Year 1 total', '~£250-350k', '~£85-100k', '£150-250k net'],
])

add_heading('Breakeven', 2)
add_bullet('Monthly costs ~£4-6k')
add_bullet('At £700/report: need ~6-9 reports/month to break even')
add_bullet('At steady state (12/week = ~50/month): ~£35k/month gross, ~£30k/month net')

doc.add_page_break()

# ============================================================
# COMPLETE TIMELINE
# ============================================================
add_heading('Complete Timeline', 1)

add_table(['Week', 'Actions', 'Dependencies'], [
    ['1', 'Form Ltd company, open bank account, appoint accountant', 'None'],
    ['2', 'Apply RICS firm registration, get PII quotes', 'Company formed'],
    ['3', 'Bind PII, register ICO, set up PAYE', 'RICS application in'],
    ['4', 'Write all policies (AML, CHP, privacy, H&S, lone worker)', 'PII bound'],
    ['5', 'Set up technology, subscriptions, equipment', 'None'],
    ['6', 'Build website, professional email, Google Business Profile', 'None'],
    ['7', 'Apply to ALL panels simultaneously', 'RICS registered, PII in place'],
    ['8', 'Submit sample reports, DBS checks', 'Panel applications in'],
    ['9-12', 'Complete panel onboarding, assessments', 'Panels processing'],
    ['10-16', 'First panel instructions arrive', 'Panel approval'],
    ['16-20', 'Volume ramps, build quality track record', 'Steady work flow'],
    ['20+', 'Steady state ~12 reports/week', 'Established'],
])

add_para('Realistic timeline: Decision to first revenue = 3-4 months', bold=True)

doc.add_page_break()

# ============================================================
# KEY REGULATORY BODIES
# ============================================================
add_heading('Key Regulatory Bodies', 1)
add_table(['Body', 'Role'], [
    ['RICS', 'Professional regulation, firm registration, AML supervision, standards'],
    ['ICO', 'Data protection registration and enforcement'],
    ['HMRC', 'Tax (PAYE, VAT, Corporation Tax)'],
    ['NCA', 'Receives SARs (suspicious activity reports)'],
    ['HSE', 'Health & Safety enforcement'],
    ['Companies House', 'Company registration and filing'],
    ['The Pensions Regulator', 'Auto-enrolment compliance'],
    ['CEDR', 'Independent dispute resolution for RICS consumer complaints'],
])

# ============================================================
# REVENUE LADDER
# ============================================================
add_heading('Revenue Ladder Strategy', 1)
add_table(['Stage', 'Source', 'Fee Range', 'Timeline'], [
    ['1', 'Panels/AMCs', '£500-700', 'Month 4+'],
    ['2', 'Specialist/bridging lenders', '£600-800', 'Month 6+'],
    ['3', 'Direct lender relationships', '£700+', 'Year 1+'],
    ['4', 'HNW landlords (portfolio)', '£700+', 'Year 1+'],
    ['5', 'High street banks direct', '£700+', 'Year 2+'],
    ['6', 'Overseas clients', '£700-1,000+', 'Year 2+'],
])

# ============================================================
# PROPVAL INTEGRATION
# ============================================================
add_heading('PropVal Integration Points', 1)
add_para('Every aspect of the firm feeds back into PropVal:')
add_bullet('600 reports/yr → SEMV training data, citation graph, UPRN spine')
add_bullet('Panel tech requirements → PropVal compliance features')
add_bullet('Lender report templates → ARTG template library')
add_bullet('SLA pressure → PropVal speed optimisation')
add_bullet('PVQ tracking → PropVal QA features')
add_bullet('Admin workflow → PropVal admin module (wife\'s feedback)')
add_bullet('Quality audits → PropVal audit trail features')

doc.add_paragraph()
p = add_para('After 1 year: Launch PropVal to market backed by 600 live, panel-accepted valuations.', bold=True)

# ============================================================
# SAVE
# ============================================================
output_path = r'C:\Users\licww\Desktop\Valuation_Firm_Setup_Guide.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
