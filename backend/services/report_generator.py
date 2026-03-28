"""Report Generator — merges structured content with template schema to produce .docx.

Takes:
  1. Structured content (metadata, valuer inputs, AI sections, comparables)
  2. Template schema (sections, branding, page setup)

Produces:
  - A professionally styled .docx file matching the template layout

Uses python-docx for .docx creation. This replaces the client-side docx-js export
when a template is selected.
"""

import io
import logging
import tempfile
from pathlib import Path

import httpx
from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)

# ── Colour constants ─────────────────────────────────────────────────────────
_ALT_ROW_COLOR = "F2F2F7"  # Light grey for alternating table rows
_WHITE = "FFFFFF"
_BORDER_COLOR = "D1D1D6"   # Subtle border grey


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert #RRGGBB to RGBColor."""
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _hex_strip(hex_color: str) -> str:
    """Strip # from hex color for XML usage."""
    return hex_color.lstrip("#")


# ── Table styling helpers ────────────────────────────────────────────────────

def _set_cell_shading(cell, hex_color: str):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell, color: str = _BORDER_COLOR, width: str = "4"):
    """Set thin borders on all sides of a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{width}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{width}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{width}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{width}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tc_borders_existing = tc_pr.find(qn('w:tcBorders'))
    if tc_borders_existing is not None:
        tc_pr.remove(tc_borders_existing)
    tc_pr.append(tc_borders)


def _set_cell_padding(cell, top=40, bottom=40, left=80, right=80):
    """Set cell padding in twips."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_mar_existing = tc_pr.find(qn('w:tcMar'))
    if tc_mar_existing is not None:
        tc_pr.remove(tc_mar_existing)
    tc_pr.append(tc_mar)


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 11,
                   color: str | None = None, font_name: str = "Calibri"):
    """Set text in a table cell with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold
    if color:
        run.font.color.rgb = _hex_to_rgb(color)


def _style_table(table, accent_color: str, font_name: str = "Calibri"):
    """Apply professional styling to an entire table: header row, alternating colours, borders."""
    accent_hex = _hex_strip(accent_color)

    for i, row in enumerate(table.rows):
        for cell in row.cells:
            _set_cell_borders(cell)
            _set_cell_padding(cell)
            if i == 0:
                # Header row: accent background, white text
                _set_cell_shading(cell, accent_hex)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.bold = True
                        r.font.name = font_name
                        r.font.size = Pt(10)
            else:
                # Alternating rows
                bg = _ALT_ROW_COLOR if i % 2 == 0 else _WHITE
                _set_cell_shading(cell, bg)


# ── Header / Footer ─────────────────────────────────────────────────────────

def _add_header_footer(doc, firm_name: str, report_ref: str, accent_color: str,
                       font_family: str):
    """Add firm name header and page number footer to the document."""
    section = doc.sections[0]

    # Header: firm name right-aligned
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(firm_name)
    run.font.size = Pt(8)
    run.font.name = font_family
    run.font.color.rgb = _hex_to_rgb(accent_color)

    # Footer: report ref left, page number right
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if report_ref:
        ref_run = fp.add_run(f"Ref: {report_ref}")
        ref_run.font.size = Pt(7)
        ref_run.font.name = font_family
        ref_run.font.color.rgb = _hex_to_rgb("#8E8E93")

    # Add page number field (right-aligned via tab stop)
    fp.add_run("\t\t")
    page_run = fp.add_run()
    page_run.font.size = Pt(7)
    page_run.font.name = font_family
    page_run.font.color.rgb = _hex_to_rgb("#8E8E93")
    # Insert PAGE field
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    page_run._r.append(fld_char_begin)
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    page_run._r.append(instr)
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    page_run._r.append(fld_char_end)


# ── Image helpers ────────────────────────────────────────────────────────────

def _download_image(url: str) -> io.BytesIO | None:
    """Download an image from URL, return BytesIO or None on failure."""
    try:
        resp = httpx.get(url, timeout=10, follow_redirects=True)
        if resp.status_code == 200 and resp.headers.get("content-type", "").startswith("image"):
            return io.BytesIO(resp.content)
    except Exception as exc:
        logger.warning("Failed to download image %s: %s", url, exc)
    return None


def _add_image(doc, url: str, width: Inches = Inches(5.5)):
    """Download and insert an image, or add placeholder text on failure."""
    img_data = _download_image(url)
    if img_data:
        doc.add_picture(img_data, width=width)
        # Centre the image
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        _add_body_text(doc, f"[Image unavailable]", "Calibri", 10, "#8E8E93")


# ── Horizontal rule ──────────────────────────────────────────────────────────

def _add_horizontal_rule(doc, color: str = _BORDER_COLOR):
    """Add a thin horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    p_bdr_existing = p_pr.find(qn('w:pBdr'))
    if p_bdr_existing is not None:
        p_pr.remove(p_bdr_existing)
    p_pr.append(p_bdr)


# ── Main generator ───────────────────────────────────────────────────────────

def generate_report_docx(
    content: dict,
    template_schema: dict,
    panel_config: dict | None = None,
) -> str:
    """Generate a .docx report from structured content and template schema.

    Args:
        content: {
            "metadata": { report_reference, report_date, client_name, ... },
            "valuer_inputs": { market_value, condition_rating, ... },
            "ai_sections": { location_description, market_commentary, ... },
            "comparables": [ { address, price, date, area, ... } ],
            "property": { address, postcode, property_type, tenure, ... },
            "firm_template": { instructions, purpose, firm_name, logo_url, ... },
            "images": { "front_elevation": "https://...", ... },
        }
        template_schema: The JSONB schema from report_templates.schema
        panel_config: Optional panel overlay config (from panel_configs.config).
            If provided, merged onto template_schema before rendering.

    Returns:
        Path to the generated .docx file (temp file).
    """
    # Apply panel overlay if provided
    if panel_config:
        from services.panel_service import merge_panel_with_template
        template_schema = merge_panel_with_template(template_schema, panel_config)
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────
    page_config = template_schema.get("page", {})
    margins = page_config.get("margins", {})
    section = doc.sections[0]
    section.page_width = Mm(210)  # A4
    section.page_height = Mm(297)

    if page_config.get("orientation") == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    # Margins (stored in twips in schema, convert to EMU: 1 twip = 635 EMU)
    if margins.get("top"):
        section.top_margin = Emu(margins["top"] * 635)
    if margins.get("bottom"):
        section.bottom_margin = Emu(margins["bottom"] * 635)
    if margins.get("left"):
        section.left_margin = Emu(margins["left"] * 635)
    if margins.get("right"):
        section.right_margin = Emu(margins["right"] * 635)

    # ── Branding ──────────────────────────────────────────────────────────
    branding = template_schema.get("branding", {})
    font_family = branding.get("font_family", "Calibri")
    font_size = branding.get("font_size", 11)
    accent_color = branding.get("accent_color", "#007AFF")

    # Set default styles
    style = doc.styles["Normal"]
    style.font.name = font_family
    style.font.size = Pt(font_size)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Style headings with accent color
    for level in range(1, 4):
        try:
            h = doc.styles[f"Heading {level}"]
            h.font.color.rgb = _hex_to_rgb(accent_color)
            h.font.name = font_family
        except KeyError:
            pass

    # ── Extract content ───────────────────────────────────────────────────
    metadata = content.get("metadata", {})
    valuer = content.get("valuer_inputs", {})
    ai_sections = content.get("ai_sections", {})
    comparables = content.get("comparables", [])
    prop = content.get("property", {})
    firm = content.get("firm_template", {})
    images = content.get("images", {})

    # ── Header & footer ───────────────────────────────────────────────────
    firm_name = firm.get("firm_name", "")
    report_ref = metadata.get("report_reference", "")
    if firm_name:
        _add_header_footer(doc, firm_name, report_ref, accent_color, font_family)

    # ── Render each section ───────────────────────────────────────────────
    first_section = True
    for tmpl_section in template_schema.get("sections", []):
        section_type = tmpl_section.get("type", "narrative")
        title = tmpl_section.get("title", "")
        page_break = tmpl_section.get("page_break_before", False)

        # Page break before non-cover sections (if configured or after cover)
        if page_break and not first_section:
            doc.add_page_break()

        # ── Cover page ────────────────────────────────────────────────────
        if section_type == "cover_page":
            first_section = False

            # Firm logo
            logo_url = firm.get("logo_url", "")
            if logo_url:
                _add_image(doc, logo_url, width=Inches(2.5))
                doc.add_paragraph()  # spacer

            # Vertical spacing
            for _ in range(3):
                doc.add_paragraph()

            # Report title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.font.size = Pt(26)
            run.font.name = font_family
            run.font.color.rgb = _hex_to_rgb(accent_color)
            run.bold = True

            _add_horizontal_rule(doc, _hex_strip(accent_color))

            # Property address
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(prop.get("address", ""))
            run.font.size = Pt(18)
            run.font.name = font_family
            run.bold = True

            # Postcode
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(prop.get("postcode", ""))
            run.font.size = Pt(14)
            run.font.name = font_family
            run.font.color.rgb = _hex_to_rgb("#636366")

            doc.add_paragraph()  # spacer

            # Metadata fields in a clean table
            meta_fields = tmpl_section.get("fields", [])
            meta_rows = []
            for field in meta_fields:
                val = (
                    metadata.get(field, "")
                    or prop.get(field, "")
                    or firm.get(field, "")
                    or ""
                )
                if val:
                    label = field.replace("_", " ").title()
                    meta_rows.append((label, str(val)))

            if meta_rows:
                table = doc.add_table(rows=0, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for label, val in meta_rows:
                    row = table.add_row()
                    _set_cell_text(row.cells[0], label, bold=True, size=10,
                                   color="#636366", font_name=font_family)
                    _set_cell_text(row.cells[1], val, size=10, font_name=font_family)
                    _set_cell_padding(row.cells[0])
                    _set_cell_padding(row.cells[1])

            # Firm name at bottom
            doc.add_paragraph()
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(firm_name)
            run.font.size = Pt(12)
            run.font.name = font_family
            run.font.color.rgb = _hex_to_rgb(accent_color)
            run.bold = True

            doc.add_page_break()

        # ── Boilerplate ───────────────────────────────────────────────────
        elif section_type == "boilerplate":
            first_section = False
            _add_heading(doc, title, accent_color, font_family)
            # Panel override_text takes precedence over firm template source_field
            text = tmpl_section.get("override_text", "")
            if not text:
                source_field = tmpl_section.get("source_field", "")
                text = firm.get(source_field, "") if source_field else ""

            subsections = tmpl_section.get("subsections", [])
            if subsections:
                for sub in subsections:
                    _add_subheading(doc, sub.get("title", ""), font_family)
                    sub_text = sub.get("override_text", "")
                    if not sub_text:
                        sub_field = sub.get("source_field", "")
                        sub_text = firm.get(sub_field, "") if sub_field else ""
                    if sub.get("type") == "data_field":
                        _render_data_fields(doc, sub.get("fields", []), valuer,
                                            font_family, accent_color)
                    elif sub_text:
                        _add_rich_text(doc, sub_text, font_family, font_size)
                    else:
                        _add_body_text(doc, "[Not configured]", font_family, font_size, "#8E8E93")
            elif text:
                _add_rich_text(doc, text, font_family, font_size)
            else:
                _add_body_text(doc, "[Not configured]", font_family, font_size, "#8E8E93")

        # ── Narrative ─────────────────────────────────────────────────────
        elif section_type == "narrative":
            first_section = False
            _add_heading(doc, title, accent_color, font_family)
            ai_key = tmpl_section.get("ai_section_key", "")
            text = ai_sections.get(ai_key, "") if ai_key else ""

            subsections = tmpl_section.get("subsections", [])
            if subsections:
                for sub in subsections:
                    _add_subheading(doc, sub.get("title", ""), font_family)
                    sub_ai_key = sub.get("ai_section_key", "")
                    sub_text = ai_sections.get(sub_ai_key, "") if sub_ai_key else ""
                    if sub.get("type") == "data_field":
                        _render_data_fields(doc, sub.get("fields", []), valuer,
                                            font_family, accent_color)
                    elif sub_text:
                        _add_rich_text(doc, sub_text, font_family, font_size)
                    else:
                        _add_body_text(doc, "[Pending]", font_family, font_size, "#FF9500")
            elif text:
                _add_rich_text(doc, text, font_family, font_size)
            else:
                _add_body_text(doc, "[Pending]", font_family, font_size, "#FF9500")

        # ── Data fields ───────────────────────────────────────────────────
        elif section_type == "data_field":
            first_section = False
            _add_heading(doc, title, accent_color, font_family)
            _render_data_fields(doc, tmpl_section.get("fields", []), valuer,
                                font_family, accent_color)

        # ── Comparables table ─────────────────────────────────────────────
        elif section_type == "comparables_table":
            first_section = False
            _add_heading(doc, title, accent_color, font_family)
            columns = tmpl_section.get("columns", ["address", "price", "date", "type", "area"])
            max_rows = tmpl_section.get("max_rows", 6)

            if comparables:
                table = doc.add_table(rows=1, cols=len(columns))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Header row
                for j, col in enumerate(columns):
                    _set_cell_text(table.rows[0].cells[j], col.replace("_", " ").title(),
                                   bold=True, size=9, font_name=font_family)

                # Data rows
                for comp in comparables[:max_rows]:
                    row = table.add_row()
                    for j, col in enumerate(columns):
                        val = comp.get(col, "")
                        if col == "price" and isinstance(val, (int, float)):
                            val = f"£{val:,.0f}"
                        _set_cell_text(row.cells[j], str(val or "—"), size=9,
                                       font_name=font_family)

                # Apply professional table styling
                _style_table(table, accent_color, font_family)
                doc.add_paragraph()  # spacing after table
            else:
                _add_body_text(doc, "No comparables adopted.", font_family, font_size, "#8E8E93")

        # ── Valuation summary ─────────────────────────────────────────────
        elif section_type == "valuation_summary":
            first_section = False
            _add_heading(doc, title, accent_color, font_family)
            subsections = tmpl_section.get("subsections", [])
            if subsections:
                for sub in subsections:
                    _add_subheading(doc, sub.get("title", ""), font_family)
                    if sub.get("type") == "narrative":
                        ai_key = sub.get("ai_section_key", "")
                        text = ai_sections.get(ai_key, "") if ai_key else ""
                        if text:
                            _add_rich_text(doc, text, font_family, font_size)
                        else:
                            _add_body_text(doc, "[Pending]", font_family, font_size, "#FF9500")
                    elif sub.get("type") == "data_field":
                        _render_data_fields(doc, sub.get("fields", []), valuer,
                                            font_family, accent_color)
            else:
                mv = valuer.get("market_value", "")
                if mv:
                    _add_horizontal_rule(doc, _hex_strip(accent_color))
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(f"£{int(float(mv)):,}" if mv else "£___________")
                    run.font.size = Pt(24)
                    run.font.bold = True
                    run.font.color.rgb = _hex_to_rgb(accent_color)
                    run.font.name = font_family
                    _add_horizontal_rule(doc, _hex_strip(accent_color))

        # ── Appendices ────────────────────────────────────────────────────
        elif section_type == "appendices":
            first_section = False
            doc.add_page_break()
            _add_heading(doc, title, accent_color, font_family)
            for sub in tmpl_section.get("subsections", []):
                _add_subheading(doc, sub.get("title", ""), font_family)
                if sub.get("type") in ("image", "image_grid"):
                    # Try to render actual images
                    img_key = sub.get("image_key", sub.get("title", "").lower().replace(" ", "_"))
                    img_url = images.get(img_key, "")
                    if img_url:
                        _add_image(doc, img_url)
                    else:
                        _add_body_text(doc, "[Image placeholder]", font_family, font_size, "#5AC8FA")
                else:
                    _add_body_text(doc, "[See attached]", font_family, font_size, "#8E8E93")

        # ── Image grid ────────────────────────────────────────────────────
        elif section_type in ("image_grid", "image"):
            first_section = False
            _add_heading(doc, title, accent_color, font_family)
            img_key = tmpl_section.get("image_key", title.lower().replace(" ", "_"))
            img_url = images.get(img_key, "")
            if img_url:
                _add_image(doc, img_url)
            else:
                _add_body_text(doc, "[Image placeholder]", font_family, font_size, "#5AC8FA")

        # ── Auto / placeholder ────────────────────────────────────────────
        elif section_type in ("auto", "placeholder"):
            first_section = False
            _add_heading(doc, title, accent_color, font_family)
            _add_body_text(doc, "[Auto-populated / placeholder]", font_family, font_size, "#8E8E93")

    # ── Save to temp file ─────────────────────────────────────────────────
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


# ── Helper functions ──────────────────────────────────────────────────────────

def _add_heading(doc, text: str, accent_color: str, font_family: str):
    """Add a styled H2 heading with bottom border."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = font_family
    run.font.color.rgb = _hex_to_rgb(accent_color)
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(16)

    # Bottom border for visual separation
    p_pr = p._p.get_or_add_pPr()
    p_bdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:color="{_hex_strip(accent_color)}"/>'
        f'</w:pBdr>'
    )
    p_bdr_existing = p_pr.find(qn('w:pBdr'))
    if p_bdr_existing is not None:
        p_pr.remove(p_bdr_existing)
    p_pr.append(p_bdr)


def _add_subheading(doc, text: str, font_family: str):
    """Add a styled H3 subheading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = font_family
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(10)


def _add_body_text(doc, text: str, font_family: str, font_size: int,
                   color: str | None = None):
    """Add body text, splitting on double newlines into paragraphs."""
    for paragraph_text in text.split("\n\n"):
        paragraph_text = paragraph_text.strip()
        if not paragraph_text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(paragraph_text)
        run.font.size = Pt(font_size)
        run.font.name = font_family
        if color:
            run.font.color.rgb = _hex_to_rgb(color)


def _add_bullet(doc, text: str, font_family: str, font_size: int,
                bold_prefix: str | None = None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = font_family
        run = p.add_run(text)
        run.font.size = Pt(font_size)
        run.font.name = font_family
    else:
        run = p.add_run(text)
        run.font.size = Pt(font_size)
        run.font.name = font_family


def _add_rich_text(doc, text: str, font_family: str, font_size: int):
    """Parse text with bullet markers and render as mixed paragraphs + bullets.

    Supports:
      - Lines starting with '- ' or '• ' → bullet points
      - Lines starting with '**text:** rest' → bullet with bold prefix
      - Everything else → normal paragraph
    """
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue

        lines = block.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect bullet lines
            if line.startswith("- ") or line.startswith("• "):
                bullet_text = line[2:].strip()

                # Check for bold prefix pattern: **Label:** rest
                if bullet_text.startswith("**") and ":**" in bullet_text:
                    idx = bullet_text.index(":**")
                    prefix = bullet_text[2:idx] + ": "
                    rest = bullet_text[idx + 3:].strip()
                    _add_bullet(doc, rest, font_family, font_size, bold_prefix=prefix)
                else:
                    _add_bullet(doc, bullet_text, font_family, font_size)
            else:
                # Normal paragraph
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                run = p.add_run(line)
                run.font.size = Pt(font_size)
                run.font.name = font_family


def _render_data_fields(doc, fields: list[str], valuer: dict, font_family: str,
                        accent_color: str = "#007AFF"):
    """Render data fields as a styled two-column table."""
    if not fields:
        return

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, field in enumerate(fields):
        val = valuer.get(field, "")
        if isinstance(val, bool):
            val = "Yes" if val else "No"
        elif val is None:
            val = ""

        row = table.add_row()
        label = field.replace("_", " ").replace("basis ", "").title()
        _set_cell_text(row.cells[0], label, bold=True, size=10,
                       color="#636366", font_name=font_family)
        _set_cell_text(row.cells[1], str(val) if val else "—", size=10,
                       font_name=font_family)

        # Alternating row shading
        bg = _ALT_ROW_COLOR if i % 2 == 0 else _WHITE
        _set_cell_shading(row.cells[0], bg)
        _set_cell_shading(row.cells[1], bg)
        _set_cell_padding(row.cells[0])
        _set_cell_padding(row.cells[1])
        _set_cell_borders(row.cells[0])
        _set_cell_borders(row.cells[1])

    doc.add_paragraph()  # spacing after table
