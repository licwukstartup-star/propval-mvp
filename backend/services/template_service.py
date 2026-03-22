"""Template parsing and AI classification service for ARTG.

Phase 2: Parse uploaded .docx files and classify sections using Claude API.

Pipeline:
  1. Extract structure from .docx (python-docx: headings, paragraphs, tables, styles)
  2. Classify sections against RICS valuation schema (Claude API)
  3. Return draft template schema JSON with confidence scores
"""

import json
import logging
import os
import tempfile
from pathlib import Path

import httpx

logger = logging.getLogger(__name__)

# ── Known RICS section types for classification ──────────────────────────────

KNOWN_SECTION_TYPES = {
    "cover_page": "Cover page with property address, dates, client name, report reference",
    "boilerplate": "Standard firm boilerplate text (instructions, disclaimers, PI insurance, etc.)",
    "narrative": "Free-text narrative written by valuer or AI (location, description, market commentary)",
    "data_field": "Structured data fields (measurements, condition, services, valuation figures)",
    "comparables_table": "Table of comparable properties with prices, dates, sizes",
    "valuation_summary": "Valuation section with market value, market rent, BIRC",
    "image_grid": "Grid of photographs",
    "image": "Single image (location plan, map)",
    "appendices": "Appendices section containing supporting documents",
    "auto": "Auto-populated from property data APIs",
    "placeholder": "Placeholder for future content (e.g., uploaded documents)",
}

RICS_SECTION_HINTS = [
    ("instructions", "boilerplate", "Instructions, scope, terms of engagement"),
    ("purpose", "boilerplate", "Purpose of valuation"),
    ("responsibility", "boilerplate", "Responsibility, third parties, reliance"),
    ("disclosure", "boilerplate", "Disclosure, conflict of interest"),
    ("basis of valuation", "data_field", "Market value, market rent, basis"),
    ("conflict", "data_field", "Conflict of interest declaration"),
    ("expertise", "boilerplate", "Valuer expertise, qualifications"),
    ("inspection", "boilerplate", "Inspection details, date, conditions"),
    ("assumptions", "data_field", "Special assumptions, caveats"),
    ("location", "narrative", "Location and locality description"),
    ("description", "narrative", "Property description, building, development"),
    ("accommodation", "narrative", "Accommodation details, rooms"),
    ("measurement", "data_field", "GIA, floor area, measurements"),
    ("site area", "data_field", "Site area, garden, land"),
    ("services", "data_field", "Gas, water, electricity, drainage"),
    ("condition", "data_field", "Condition rating, notes, defects"),
    ("environmental", "boilerplate", "Environmental matters, contamination"),
    ("asbestos", "boilerplate", "Asbestos disclaimer"),
    ("flood", "auto", "Flood risk assessment"),
    ("fire", "boilerplate", "Fire risk, cladding, EWS"),
    ("tenure", "auto", "Tenure details, leasehold, freehold"),
    ("market", "narrative", "Market commentary, trends, demand"),
    ("comparable", "comparables_table", "Comparable evidence, transactions"),
    ("valuation", "valuation_summary", "Valuation figure, opinion of value"),
    ("methodology", "boilerplate", "Valuation methodology"),
    ("market rent", "data_field", "Market rent assessment"),
    ("market value", "data_field", "Market value figure"),
    ("suitable security", "data_field", "Mortgage security suitability"),
    ("birc", "data_field", "Buildings insurance reinstatement cost"),
    ("photograph", "image_grid", "Photographs, images"),
    ("appendix", "appendices", "Appendices, supporting documents"),
    ("epc", "placeholder", "EPC certificate"),
    ("terms", "placeholder", "Terms of engagement"),
    ("location plan", "image", "Location plan, map"),
]


# ── Step 1: Parse .docx structure ────────────────────────────────────────────

async def parse_docx_structure(file_bytes: bytes) -> dict:
    """Extract document structure from .docx bytes using python-docx.

    Returns a dict with:
      - sections: list of { title, level, paragraphs, has_table, style }
      - metadata: page size, margins, fonts detected
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    # Write to temp file (python-docx needs a file path or file-like object)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        doc = DocxDocument(tmp_path)

        # Extract page layout
        section = doc.sections[0] if doc.sections else None
        page_info = {}
        if section:
            page_info = {
                "width_mm": round(section.page_width.mm, 1) if section.page_width else None,
                "height_mm": round(section.page_height.mm, 1) if section.page_height else None,
                "margin_top_mm": round(section.top_margin.mm, 1) if section.top_margin else None,
                "margin_bottom_mm": round(section.bottom_margin.mm, 1) if section.bottom_margin else None,
                "margin_left_mm": round(section.left_margin.mm, 1) if section.left_margin else None,
                "margin_right_mm": round(section.right_margin.mm, 1) if section.right_margin else None,
            }

        # Detect the most common body font size (to identify larger headings)
        body_font_sizes: list[float] = []
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size:
                    body_font_sizes.append(run.font.size.pt)
        typical_body_size = max(set(body_font_sizes), key=body_font_sizes.count) if body_font_sizes else 11

        # Walk paragraphs and detect structure
        extracted_sections: list[dict] = []
        current_section: dict | None = None

        import re
        numbered_pattern = re.compile(r"^\d+(\.\d+)*\.?\s+\S")  # "1.0 ...", "2.1 ...", "3 ..."

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            if not text:
                continue

            # ── Detect headings via multiple signals ──
            is_heading = False
            level = 1

            # Signal 1: Word heading styles (strongest signal)
            if style_name.startswith("Heading") or "heading" in style_name.lower():
                is_heading = True
                if "2" in style_name:
                    level = 2
                elif "3" in style_name:
                    level = 3
                elif "4" in style_name:
                    level = 4

            # Signal 2: Bold text, short enough to be a heading
            elif para.runs and len(text) < 150:
                all_bold = all(r.bold for r in para.runs if r.text.strip())
                if all_bold and len(para.runs) > 0:
                    is_heading = True

            # Signal 3: ALL CAPS text (common in valuation reports)
            if not is_heading and text == text.upper() and len(text) > 3 and len(text) < 120 and any(c.isalpha() for c in text):
                is_heading = True

            # Signal 4: Numbered section pattern ("1.0 Instructions", "2.1 Location")
            if not is_heading and numbered_pattern.match(text) and len(text) < 150:
                # Check if it looks like a section header (not just "1. The property is...")
                words = text.split()
                if len(words) <= 10:
                    is_heading = True
                    # Infer level from numbering depth: "1" = level 1, "1.1" = level 2, "1.1.1" = level 3
                    num_part = text.split()[0].rstrip(".")
                    dots = num_part.count(".")
                    level = min(dots + 1, 4)

            # Signal 5: Font size significantly larger than body text
            if not is_heading and para.runs and len(text) < 150:
                run_size = para.runs[0].font.size
                if run_size and run_size.pt >= typical_body_size + 2:
                    is_heading = True

            # Signal 6: Underlined short text (used as headers in some reports)
            if not is_heading and para.runs and len(text) < 120:
                all_underlined = all(r.underline for r in para.runs if r.text.strip())
                if all_underlined and len(text) > 3:
                    is_heading = True
                    level = 2

            if is_heading:
                # Save previous section
                if current_section:
                    extracted_sections.append(current_section)

                current_section = {
                    "title": text,
                    "level": level,
                    "paragraphs": [],
                    "has_table": False,
                    "style": style_name,
                    "first_paragraph": "",
                }
            elif current_section:
                current_section["paragraphs"].append(text)
                if not current_section["first_paragraph"]:
                    current_section["first_paragraph"] = text[:200]

        # Don't forget last section
        if current_section:
            extracted_sections.append(current_section)

        # Check for tables in document
        table_count = len(doc.tables)

        # Detect fonts used
        fonts_seen = set()
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    fonts_seen.add(run.font.name)

        return {
            "sections": extracted_sections,
            "table_count": table_count,
            "page": page_info,
            "fonts": list(fonts_seen)[:10],
            "total_paragraphs": len(doc.paragraphs),
        }

    finally:
        Path(tmp_path).unlink(missing_ok=True)


# ── Step 2: Classify sections with Claude API ────────────────────────────────

async def classify_sections_with_ai(parsed_structure: dict) -> dict:
    """Use Claude API to classify each section against known RICS types.

    Returns a template schema draft with confidence scores.
    """
    api_key = os.getenv("ANTHROPIC_API_KEY") or os.getenv("OPENAI_API_KEY", "")

    if not api_key:
        # Fallback: rule-based classification
        logger.warning("No AI API key configured — using rule-based classification")
        return _rule_based_classify(parsed_structure)

    # Build section summaries for the prompt
    section_summaries = []
    for i, sec in enumerate(parsed_structure.get("sections", [])):
        summary = f"{i+1}. \"{sec['title']}\" (level {sec['level']})"
        if sec.get("first_paragraph"):
            summary += f" — starts with: \"{sec['first_paragraph'][:100]}...\""
        summary += f" — {len(sec.get('paragraphs', []))} paragraphs"
        section_summaries.append(summary)

    prompt = f"""You are analysing a UK residential property valuation report template.
The document has {len(section_summaries)} sections. Classify each section into one of these types:

Types: {', '.join(KNOWN_SECTION_TYPES.keys())}

Type descriptions:
{chr(10).join(f'- {k}: {v}' for k, v in KNOWN_SECTION_TYPES.items())}

Sections found in document:
{chr(10).join(section_summaries)}

For each section, respond with a JSON array where each item has:
- "index": section number (1-based)
- "title": the section title
- "type": one of the types above
- "confidence": 0.0 to 1.0
- "ai_section_key": if type is "narrative", suggest which AI key to use (one of: location_description, subject_development, subject_building, subject_property, market_commentary, valuation_considerations) or null
- "source_field": if type is "boilerplate", suggest which firm template field to use (one of: instructions, purpose, responsibility, disclosure, pi_insurance, expertise, inspection, environmental, asbestos, fire_risk, methodology, general_comments) or null

Respond ONLY with the JSON array, no other text."""

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": api_key,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                },
                json={
                    "model": "claude-haiku-4-5-20251001",
                    "max_tokens": 4096,
                    "messages": [{"role": "user", "content": prompt}],
                },
            )
            resp.raise_for_status()
            data = resp.json()

            # Extract text content
            text = data.get("content", [{}])[0].get("text", "")

            # Parse JSON from response
            # Strip markdown code fences if present
            text = text.strip()
            if text.startswith("```"):
                text = text.split("\n", 1)[1] if "\n" in text else text[3:]
            if text.endswith("```"):
                text = text[:-3]
            text = text.strip()
            if text.startswith("json"):
                text = text[4:].strip()

            classifications = json.loads(text)
            return _build_schema_from_classifications(parsed_structure, classifications)

    except Exception as exc:
        logger.error("AI classification failed: %s — falling back to rules", exc)
        return _rule_based_classify(parsed_structure)


# ── Rule-based fallback ──────────────────────────────────────────────────────

def _rule_based_classify(parsed_structure: dict) -> dict:
    """Simple keyword-based classification when no AI is available."""
    sections = parsed_structure.get("sections", [])
    classified: list[dict] = []

    for i, sec in enumerate(sections):
        title_lower = sec["title"].lower()
        matched_type = "narrative"  # default
        confidence = 0.4
        ai_key = None
        source_field = None

        for hint_keyword, hint_type, _desc in RICS_SECTION_HINTS:
            if hint_keyword in title_lower:
                matched_type = hint_type
                confidence = 0.7
                # Set AI key for narrative sections
                if hint_type == "narrative":
                    if "location" in hint_keyword:
                        ai_key = "location_description"
                    elif "description" in hint_keyword or "building" in hint_keyword:
                        ai_key = "subject_building"
                    elif "accommodation" in hint_keyword:
                        ai_key = "subject_property"
                    elif "market" in hint_keyword:
                        ai_key = "market_commentary"
                    elif "valuation" in hint_keyword and "consideration" in title_lower:
                        ai_key = "valuation_considerations"
                elif hint_type == "boilerplate":
                    source_field = hint_keyword
                break

        classified.append({
            "index": i + 1,
            "title": sec["title"],
            "type": matched_type,
            "confidence": confidence,
            "ai_section_key": ai_key,
            "source_field": source_field,
        })

    return _build_schema_from_classifications(parsed_structure, classified)


def _build_schema_from_classifications(parsed_structure: dict, classifications: list[dict]) -> dict:
    """Convert classifications into a template schema."""
    page = parsed_structure.get("page", {})
    fonts = parsed_structure.get("fonts", [])

    schema_sections = []
    for cls in classifications:
        section: dict = {
            "id": f"section_{cls['index']}",
            "type": cls["type"],
            "title": cls["title"],
        }
        if cls.get("ai_section_key"):
            section["ai_section_key"] = cls["ai_section_key"]
        if cls.get("source_field"):
            section["source_field"] = cls["source_field"]
        if cls["type"] == "comparables_table":
            section["columns"] = ["address", "price", "date", "type", "area", "price_per_sqm"]
            section["max_rows"] = 6
        if cls["type"] == "image_grid":
            section["layout"] = "2x3"
        if cls["type"] == "cover_page":
            section["fields"] = ["property_address", "valuation_date", "client_name", "report_ref"]

        schema_sections.append(section)

    return {
        "schema": {
            "version": "1.0",
            "page": {
                "size": "A4",
                "margins": {
                    "top": int((page.get("margin_top_mm") or 25.4) / 25.4 * 1440),
                    "right": int((page.get("margin_right_mm") or 25.4) / 25.4 * 1440),
                    "bottom": int((page.get("margin_bottom_mm") or 25.4) / 25.4 * 1440),
                    "left": int((page.get("margin_left_mm") or 25.4) / 25.4 * 1440),
                },
                "orientation": "portrait",
            },
            "branding": {
                "font_family": fonts[0] if fonts else "Calibri",
                "font_size": 11,
                "accent_color": "#007AFF",
            },
            "header": {"layout": "logo-left-text-right", "content": ["{{firm_name}}", "{{firm_address}}"]},
            "footer": {"content": "Page {{page_number}} of {{total_pages}}"},
            "sections": schema_sections,
        },
        "classifications": classifications,
        "metadata": {
            "fonts_detected": fonts,
            "total_sections": len(classifications),
            "table_count": parsed_structure.get("table_count", 0),
        },
    }
