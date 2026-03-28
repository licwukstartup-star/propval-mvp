from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return h

def add_bold_para(bold_text, normal_text=''):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p

def add_bullet(text, suffix=''):
    p = doc.add_paragraph(style='List Bullet')
    if suffix:
        run = p.add_run(text)
        run.bold = True
        p.add_run(suffix)
    else:
        p.add_run(text)
    return p

def add_mono(text):
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for r, row_data in enumerate(rows):
        for c, text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = text
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    return table

# ========== TITLE ==========
title = doc.add_heading('Vibe Valuation', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0xFF, 0x7F, 0x00)
subtitle = doc.add_paragraph('The IDE for Valuations')
subtitle.style = doc.styles['Subtitle']

# ========== CONTEXT ==========
add_heading('Context', level=1)

p = doc.add_paragraph()
run = p.add_run('The insight: ')
run.bold = True
p.add_run('Programming went from terminal \u2192 IDE + AI copilot. Valuation is still stuck in the terminal era.')

add_table(
    ['Programmer today', 'Valuer today (terminal era)', 'Valuer in PropVal (IDE era)'],
    [
        ['VS Code = integrated workspace', 'Paper + Word + Excel in separate windows', 'Single PropVal workspace'],
        ['Git = version history', 'No history, overwrite files', 'Immutable report copies + audit trail'],
        ['GitHub Copilot = ghost text', 'Type everything manually', 'Inline suggestions from real property data'],
        ['Claude Code = chat + commands', 'No AI at all', 'Conversational sidebar: "add flood risk here"'],
        ['MCP = connected tools', 'Manual copy-paste between apps', 'Auto-enrichment from 19+ APIs'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('What this solves: ')
run.bold = True
p.add_run('A valuer currently spends 2-4 hours per report manually typing text that is 80% derivable from data they already have. Vibe Valuation turns that into a 15-minute refinement loop \u2014 AI drafts from real data, valuer accepts/edits/refines.')

# ========== REPORT ANATOMY ==========
add_heading('Report Anatomy \u2014 What the Valuer Actually Does', level=1)

p = doc.add_paragraph()
p.add_run('Analysis of real Grant Stanley RICS reports (house + flat) reveals:')

add_table(
    ['Report content', '% of report', 'Source', 'Vibe Valuation role'],
    [
        ['Boilerplate (T&Cs, standards, disclaimers)', '~40%', 'Template', 'Auto-fill from ARTG'],
        ['Data lookups (EPC, flood, council tax)', '~15%', 'APIs', 'Auto-populate from enrichment'],
        ['Location description', '~8%', 'Area data', 'AI drafts, valuer reviews'],
        ['Market commentary', '~8%', 'RICS survey', 'Auto-insert'],
        ['Property description + photos', '~12%', 'Inspection', 'Valuer inputs, AI assists phrasing'],
        ['Comparables + valuation reasoning', '~10%', 'Judgment', 'Valuer selects, AI drafts narrative'],
        ['Value opinion + measurement', '~7%', 'Judgment', 'Valuer decides, system records'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("The valuer's irreplaceable contribution is ~25-30%. ")
run.bold = True
p.add_run('Vibe Valuation automates the other 70-75%.')

# ========== PROVEN PATTERNS ==========
add_heading('Proven Patterns \u2014 Borrowed from Industry Leaders', level=1)

# Pattern 1
add_heading('Pattern 1: Ghost Text Inline Suggestions (from GitHub Copilot)', level=2)
add_bullet('Dimmed ghost text appears at cursor after typing pause')
add_bullet('Tab', ' = accept all, Ctrl+Right = accept next word, Esc = dismiss')
add_bullet('Next Edit Suggestions (NES) predict where to edit next, not just what to type')
add_bullet('Borrow:', ' Same UX in TipTap \u2014 ghost text, Tab accept, partial accept per word/sentence')

# Pattern 2
add_heading('Pattern 2: Ambient Capture \u2192 Structured Draft (from DAX Copilot + Nabla)', level=2)
add_bullet('Doctor talks to patient normally, AI listens and transcribes')
add_bullet('Generates structured clinical note in EHR format in <20 seconds')
add_bullet('Doctor reviews draft, edits, signs off \u2014 saves 5 min per encounter')
add_bullet('DAX embedded inside Epic; Nabla as Chrome extension \u2014 neither is a separate app')
add_bullet('Borrow (Phase 2):', ' Valuer speaks observations during inspection \u2192 AI generates Property Description in RICS format')

# Pattern 3
add_heading('Pattern 3: Playbook-Driven Workflow (from Harvey AI)', level=2)
add_bullet('Firm uploads commercial playbook (standard positions, fallbacks, preferred language)')
add_bullet('Harvey applies playbook to incoming contracts, generates redlines + comment bubbles')
add_bullet('Workflow Builder: no-code tool to design custom agents from proprietary knowledge')
add_bullet('Borrow:', " ARTG template IS the playbook. Firm's report template defines structure, tone, boilerplate. Already built \u2014 validates ARTG architecture.")

# Pattern 4
add_heading('Pattern 4: Inline Accept/Reject with Track Changes (from Spellbook)', level=2)
add_bullet('Suggestions appear as inline redlines inside Word \u2014 not a separate panel')
add_bullet('Lawyer accepts, rejects, or edits each suggestion clause by clause')
add_bullet('"Watches as you draft and redline, learns from your edits, adapts in real time"')
add_bullet('Borrow:', ' Show AI suggestions as distinguishable insertions. Every edit feeds valuer_feedback \u2192 system adapts to valuer style.')

# Pattern 5
add_heading('Pattern 5: Work Inside Existing Tools (universal)', level=2)
add_bullet('Spellbook \u2192 inside Word. DAX \u2192 inside Epic. Harvey \u2192 Word + M365. Copilot \u2192 VS Code.')
add_bullet('No successful professional copilot is a standalone app.', '')
add_bullet('Borrow:', ' Vibe Valuation lives inside the TipTap editor. Not a separate tab. The editor IS the workspace.')

# Pattern 6
add_heading('Pattern 6: Review-Not-Generate Principle (universal)', level=2)
add_bullet('AI generates draft. Professional reviews, edits, takes responsibility.')
add_bullet('DAX: "draft clinical note" \u2192 doctor signs. Spellbook: "suggested redlines" \u2192 lawyer accepts.')
add_bullet('Borrow:', " Vibe Valuation never auto-fills the report. Every AI suggestion requires explicit valuer acceptance. The valuer's signature is on the report.")

# ========== THE DIFFERENTIATOR ==========
add_heading('The Differentiator Nobody Has', level=1)

p = doc.add_paragraph()
p.add_run('Every professional copilot works with ')
run = p.add_run('one type of context:')
run.bold = True

add_table(
    ['Tool', 'Context it knows'],
    [
        ['Harvey', 'The contract text'],
        ['DAX / Nabla', 'The doctor-patient conversation'],
        ['GitHub Copilot', 'The codebase'],
        ['Spellbook', 'The contract + firm playbook'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Vibe Valuation knows ALL of these simultaneously:')
run.bold = True

add_bullet('Property data (80+ fields from 19 APIs)')
add_bullet('Comparables (selected, analysed, with \u00a3/sqft)')
add_bullet('Valuation model output (SEMV distribution)')
add_bullet('Market data (HPI trends, RICS survey)')
add_bullet('Template (ARTG structure + firm boilerplate)')
add_bullet('What the valuer is currently typing')
add_bullet('What the valuer has already written in other sections')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('No existing professional copilot has this depth of structured domain data feeding its suggestions. ')
run.bold = True
p.add_run('The UX patterns are borrowed. The context richness is unique.')

# ========== PRODUCT VISION ==========
add_heading('Product Vision \u2014 Three Layers', level=1)

add_heading('Layer 1: Inline Copilot (ghost text) \u2014 Phase 1', level=2)
add_bullet('Valuer types in the TipTap editor')
add_bullet('After ~1 second pause, AI suggests the next sentence as grey ghost text')
add_bullet('Context-aware: knows which section, all property data, comps, SEMV')
add_bullet('Tab', ' = accept all, Ctrl+Right = accept next word (partial accept, from Copilot)')
add_bullet('Ctrl+Space', ' = explicitly request suggestion')
add_bullet('Pause-based mode toggleable')
add_bullet('Every accept/reject/edit feeds valuer_feedback \u2192 training flywheel (from Spellbook)')

doc.add_paragraph()
add_heading('Layer 2: Conversational Sidebar \u2014 Phase 1', level=2)
add_bullet('Current sidebar transforms from "click Generate" into chat (from Harvey Edit Mode)')
add_bullet('Valuer types: "make this more formal", "add the lease details", "shorter"')
add_bullet('AI sees: cursor position, selected text, full property context, what\'s already written')
add_bullet('Responses insert/replace text in editor directly')
add_bullet('Keep section-generation as "quick actions" above chat')

doc.add_paragraph()
add_heading('Layer 3: Ambient Voice \u2192 Structured Draft \u2014 Phase 2', level=2)
add_bullet('Valuer speaks observations during/after inspection (from DAX/Nabla pattern)')
add_bullet('AI generates Property Description + Accommodation Schedule in RICS format')
add_bullet('Valuer reviews, edits, signs off')
add_bullet('Hybrid mode: ambient + dictation + typing (from Nabla flexibility)')

# ========== WHY THIS WINS ==========
add_heading('Why This Wins', level=1)
add_bullet('No competitor has this.', ' Proptech has AVMs and template generators \u2014 none have an inline copilot with structured domain context.')
add_bullet('Training flywheel.', ' 20 valuers \u00d7 400 reports/month = 8,000 feedback signals/month. System learns each valuer\'s style (from Spellbook).')
add_bullet('The 15-minute report becomes real.', ' ARTG (playbook, from Harvey) + Vibe Valuation (copilot) + auto-enrichment (APIs) = only 25% left for the valuer.')
add_bullet('Scales with the platform.', ' Same copilot for 20 or 1,000 valuers. More valuers = better suggestions.')
add_bullet('Proven patterns, unique context.', ' UX borrowed from $2B+ companies. Domain data depth is ours alone.')

# ========== BUILDABLE SCOPE ==========
add_heading('Buildable Scope \u2014 Phase 1 MVP', level=1)

add_heading('1. TipTap Inline Suggestion Extension', level=2)
add_bullet('Custom TipTap extension rendering ghost text (dimmed, grey) after cursor')
add_bullet('Typing pause (~1s debounce) OR Ctrl+Space hotkey triggers suggestion')
add_bullet('Tab = accept all, Ctrl+Right = accept next word (partial accept)')
add_bullet('Esc or any other key = dismiss')
add_bullet('Toggle in toolbar: "AI Assist: ON/OFF"')

add_heading('2. Context-Aware Suggestion Endpoint', level=2)
add_bullet('POST /api/ai-suggest')
add_bullet('Input: section_key, text before/after cursor, property_data, comparables, semv_output')
add_bullet('Output: suggestion string (1-2 sentences)')
add_bullet('Existing AI fallback chain (Groq \u2192 Cerebras \u2192 Gemini)')
add_bullet('Prompt tuned per section type via prompt_registry')

add_heading('3. Sidebar Chat Evolution', level=2)
add_bullet('Replace "Generate" buttons with text input + message history')
add_bullet('Valuer types instruction \u2192 AI generates/modifies text in context')
add_bullet('"Insert at cursor" / "Replace selection" buttons on each response')
add_bullet('Keep section-generation as "quick actions" above chat')

add_heading('4. Feedback Capture (Training Flywheel)', level=2)
add_bullet('Tab (accept): log { section_key, suggestion, action: "accepted" }')
add_bullet('Dismiss: log { section_key, suggestion, action: "dismissed" }')
add_bullet('Edit after accept: log { section_key, suggestion, valuer_edit, action: "edited" }')
add_bullet('All \u2192 valuer_feedback table (already exists)')

add_heading('What NOT to Build in Phase 1', level=2)
add_bullet('Voice/ambient input (Phase 2 \u2014 DAX/Nabla pattern)')
add_bullet('Per-firm style learning (Phase 2 \u2014 DAX configurable styles)')
add_bullet('Slash commands in editor (Phase 2)')
add_bullet('Fine-tuned model (use general LLM + good prompts first)')

# ========== ARCHITECTURE ==========
add_heading('Architecture Sketch', level=1)

add_mono("""Valuer types in TipTap Editor
        \u2502
        \u25bc (1s pause or Ctrl+Space)
InlineSuggestion Extension
        \u2502
        \u25bc gathers context
{ section_key, text_before, text_after, property_data, comps, semv }
        \u2502
        \u25bc POST /api/ai-suggest
Backend: generate_suggestion()  (Groq \u2192 Cerebras \u2192 Gemini)
        \u2502
        \u25bc
Ghost text rendered after cursor (dimmed grey)
        \u2502
    \u250c\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
  [Tab] [Ctrl+\u2192] [Dismiss]
 Accept  Accept   Ignore
   all    word
    \u2502      \u2502        \u2502
    \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
              \u25bc
    valuer_feedback table (training flywheel)""")

doc.add_paragraph()
add_bold_para('Sidebar chat flow:')

add_mono("""Valuer types in chat: "add flood risk"
        \u2502
        \u25bc
Backend: generate_chat_response()
  context = { selected_text, cursor_section, full_property_data, chat_history }
        \u2502
        \u25bc
AI response shown in sidebar
        \u2502
  [Insert at cursor] / [Replace selection]
        \u2502
        \u25bc
Text inserted into TipTap editor \u2192 valuer_feedback table""")

# ========== VERIFICATION ==========
add_heading('How to Test', level=1)
add_bullet('Ghost text: Start typing, pause \u2014 grey suggestion appears with property-specific data')
add_bullet('Tab accept: Press Tab \u2192 suggestion becomes real text')
add_bullet('Partial accept: Press Ctrl+Right \u2192 only next word accepted (from Copilot)')
add_bullet('Toggle: Turn off AI Assist \u2192 no suggestions on pause')
add_bullet('Chat sidebar: Type "describe the location" \u2192 AI responds with real property data')
add_bullet('Insert: Click "Insert at cursor" \u2192 text appears at correct position')
add_bullet('Feedback: Check valuer_feedback table \u2192 accepted/dismissed/edited logged')
add_bullet('Performance: Suggestion within 1-2 seconds of pause')

# ========== COMPETITIVE POSITIONING ==========
add_heading('Competitive Positioning', level=1)

p = doc.add_paragraph()
p.style = doc.styles['Intense Quote']
p.add_run("\"Every valuation firm types the same 80% of every report from scratch. PropVal's Vibe Valuation is the first inline AI copilot for property valuation \u2014 it knows your property data, your comparables, your valuation, and your house style. The valuer's job shifts from typing to refining. That's how you go from 3-hour reports to 15-minute reports.\"")

doc.add_paragraph()
p = doc.add_paragraph()
p.style = doc.styles['Intense Quote']
p.add_run('"Harvey did this for lawyers. DAX did this for doctors. Vibe Valuation does this for valuers \u2014 with richer structured context than any of them."')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('ARTG gives you the skeleton. ')
run = p.add_run('Vibe Valuation fills in the muscle.')
run.bold = True

# ========== RESEARCH SOURCES ==========
add_heading('Research Sources', level=1)
add_bullet('Harvey AI', ' \u2014 $2B+ legal AI: playbook-driven workflow, Word integration, redlining')
add_bullet('Spellbook', ' \u2014 Contract AI inside Word: inline suggestions, learns from edits')
add_bullet('DAX Copilot', ' (Microsoft/Nuance) \u2014 Ambient clinical AI: listen \u2192 structured note, 400+ health systems')
add_bullet('Nabla', ' \u2014 Medical AI scribe: ambient + dictation hybrid, $70M raised')
add_bullet('Clio', ' \u2014 #1 legal platform: AI across draft, manage, bill')
add_bullet('GitHub Copilot', ' \u2014 Ghost text, Tab accept, partial accept, NES, 100M+ users')

output_path = r'C:\Users\licww\Desktop\propval-mvp\Vibe Valuation\Vibe Valuation Plan.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
