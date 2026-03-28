from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def h(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return heading

def p(text='', bold_prefix=''):
    para = doc.add_paragraph()
    if bold_prefix:
        run = para.add_run(bold_prefix)
        run.bold = True
        para.add_run(text)
    else:
        para.add_run(text)
    return para

def b(text, suffix=''):
    para = doc.add_paragraph(style='List Bullet')
    if suffix:
        run = para.add_run(text)
        run.bold = True
        para.add_run(suffix)
    else:
        para.add_run(text)
    return para

def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(hdr)
        run.bold = True
        run.font.size = Pt(9)
    for r, row_data in enumerate(rows):
        for c, text in enumerate(row_data):
            cell = t.rows[r + 1].cells[c]
            cell.text = text
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    return t

def mono(text):
    para = doc.add_paragraph()
    para.style = doc.styles['No Spacing']
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return para

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('Vibe Valuation', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0xFF, 0x7F, 0x00)
sub = doc.add_paragraph('User Guide')
sub.style = doc.styles['Subtitle']

# ============================================================
# WHAT IS VIBE VALUATION
# ============================================================
h('What is Vibe Valuation?', 1)
p("Vibe Valuation is PropVal's built-in AI writing assistant for valuation reports. It works like GitHub Copilot for code \u2014 but for RICS-compliant property valuation reports. Instead of typing every sentence from scratch, you type a few words and the AI suggests the rest, using your actual property data, comparables, and valuation figures.")
p('', 'Two tools, one editor:')
b('Ghost Text', ' \u2014 grey suggestion text appears as you type. Press Tab to accept.')
b('AI Chat', ' \u2014 a conversational sidebar where you can ask the AI to draft, rephrase, or edit text.')
p('Both tools know your property data. They never invent facts. You always have final control.')

# ============================================================
# GETTING STARTED
# ============================================================
h('Getting Started', 1)
h('1. Open the Report Editor', 2)
b('Search for a property by postcode')
b('Select the address')
b('Navigate to the Report Typing tab')
b('Select Editor view')
p('You will see the TipTap document editor with a toolbar at the top and the AI sidebar on the right.')

h('2. Check AI Assist is ON', 2)
p('Look at the toolbar \u2014 you should see a button that says "Assist ON" with a sparkle icon. If it says "Assist OFF", click it to turn on ghost text suggestions.')
p('', 'When AI Assist is ON: ')
b('The button shows green with "Assist ON"')
b('Ghost text will appear after you pause typing')
b('Ctrl+Space triggers a suggestion immediately')
p('', 'When AI Assist is OFF: ')
b('The button shows grey with "Assist OFF"')
b('No ghost text appears')
b('The AI Chat sidebar still works independently')

# ============================================================
# GHOST TEXT
# ============================================================
h('Ghost Text (Inline Suggestions)', 1)
h('How It Works', 2)
b('Type normally', ' in any section of the report')
b('Pause for 1 second', ' \u2014 grey italic text appears after your cursor')
b('The suggestion continues your sentence using real property data')

h('Keyboard Shortcuts', 2)
table(
    ['Action', 'Shortcut', 'What happens'],
    [
        ['Accept full suggestion', 'Tab', 'The entire grey text becomes part of your report'],
        ['Accept next word', 'Ctrl + \u2192 (Right Arrow)', 'Only the next word is accepted, rest stays as ghost text'],
        ['Dismiss suggestion', 'Esc', 'Ghost text disappears'],
        ['Request suggestion now', 'Ctrl + Space', 'Triggers a suggestion immediately without waiting 1 second'],
        ['Keep typing', 'Any letter key', 'Ghost text disappears and your typing continues normally'],
    ]
)

h('Tips for Best Results', 2)
b('Write the start of a sentence', ' \u2014 the AI works best when it can see where you\'re heading. "The property is located in" gives better results than just "The".')
b('Be in the right section', ' \u2014 the AI detects which report section you\'re in (Location, Property Description, Valuation, etc.) and adjusts its suggestions accordingly.')
b('Use Ctrl+Space in empty sections', ' \u2014 if you\'ve just started a new section and want a first sentence, press Ctrl+Space to request a suggestion.')
b('Partial accept is powerful', ' \u2014 use Ctrl+Right Arrow to accept one word at a time. This lets you take the parts you like and then continue typing your own words.')
b('Don\'t fight it', ' \u2014 if the suggestion isn\'t right, just keep typing. The ghost text disappears instantly and won\'t interfere.')

h('What the AI Knows', 2)
p('Ghost text suggestions are informed by:')
b('Property address, postcode, and borough')
b('Property type, built form, and construction era')
b('Floor area, number of rooms, EPC rating')
b('Tenure, heating type, council tax band')
b('Flood risk levels (rivers/sea and surface water)')
b('Your adopted comparables (addresses and prices)')
b('SEMV valuation output (mean, confidence interval)')
b('What you\'ve already written in the current section')
p('The AI never invents addresses, prices, measurements, or dates. If it doesn\'t have the data, it won\'t guess.')

# ============================================================
# AI CHAT
# ============================================================
h('AI Chat Sidebar', 1)
h('Opening the Chat', 2)
b('Click the "AI" button at the right end of the toolbar')
b('The sidebar opens \u2014 you\'ll see two tabs: Chat and Sections')
b('Chat', ' is the conversational assistant (default tab)')
b('Sections', ' contains the original section generators')

h('How to Use the Chat', 2)
p('Type a natural language instruction in the text box at the bottom and press Enter.')
p('', 'Example instructions:')
table(
    ['What you type', 'What the AI does'],
    [
        ['"describe the location"', 'Drafts a location description using the property\'s postcode, borough, and nearby amenities'],
        ['"add the flood risk details"', 'Writes flood risk text using the Environment Agency data for this property'],
        ['"make this more formal"', 'Rephrases your selected text in formal RICS register'],
        ['"shorten this paragraph"', 'Condenses the selected text while keeping key facts'],
        ['"add the lease details"', 'Writes tenure text using the property\'s lease data'],
        ['"draft valuation considerations"', 'Writes a structured comparable analysis using your adopted comps'],
        ['"rewrite without marketing language"', 'Removes subjective adjectives and makes the text factual'],
    ]
)

h('Working with Selected Text', 2)
b('Select text', ' in the editor (highlight it with your mouse or Shift+Arrow keys)')
b('The sidebar shows a blue "Selected:" indicator at the top')
b('Type your instruction \u2014 the AI knows what you\'ve selected')
b('When the AI responds, you\'ll see action buttons:')
p('   \u2022  "Insert at cursor" \u2014 adds the text at your current cursor position')
p('   \u2022  "Replace selection" \u2014 replaces your highlighted text with the AI\'s version')
p('   \u2022  "Copy" \u2014 copies the response to your clipboard')

h('Chat Tips', 2)
b('Be specific', ' \u2014 "add flood risk" works better than "add some details"')
b('Iterate', ' \u2014 you can follow up: "now make it shorter", "add the surface water risk too"')
b('The chat remembers', ' \u2014 it keeps your last few messages as context')
b('Shift+Enter', ' for multi-line instructions')
b('Click the suggestion hints', ' when the chat is empty \u2014 they pre-fill common instructions')

h('Sections Tab (Quick Actions)', 2)
p('Switch to the Sections tab for the original full-section generators:')
b('Click a section (e.g., "2.2 Location Description")')
b('Click "Generate" \u2014 the AI writes the full section text')
b('Preview it in the expandable panel')
b('Click "Insert" to place it into the document at the correct section')
b('Edit in the document as needed')
p('These are best for generating entire sections from scratch. The Chat tab is better for refining, rephrasing, or adding specific details.')

# ============================================================
# WORKFLOW
# ============================================================
h('Workflow \u2014 Writing a Report with Vibe Valuation', 1)

h('Step 1: Let the template do the heavy lifting', 2)
p('When you load a report template, PropVal automatically fills all boilerplate sections (T&Cs, disclaimers, valuation standards), API-populated data (EPC, flood risk, council tax, transport links), and case metadata (dates, client name, reference numbers).')
p('', 'You don\'t need to type any of this. ')
p('It\'s already in the document.')

h('Step 2: Generate narrative sections', 2)
p('Use the Sections tab to generate first drafts of:')
b('Location Description (Section 2.2)')
b('Market Commentary (Section 3.3)')
b('Valuation Considerations (Section 3.6)')
p('Click Generate \u2192 review \u2192 Insert. Then edit the text in the document.')

h('Step 3: Type with ghost text assistance', 2)
p('For sections that need your professional input:')
b('Property Description', ' \u2014 start typing what you observed during inspection. Ghost text will suggest continuations using the property data.')
b('Condition', ' \u2014 type your observations. Ghost text completes with appropriate surveyor language.')
b('Accommodation schedule', ' \u2014 type room names. Ghost text suggests descriptions based on property type and era.')
p('Use Tab to accept good suggestions. Use Ctrl+Right to accept word-by-word. Keep typing to ignore.')

h('Step 4: Refine with the Chat', 2)
p('After your first pass:')
b('Select a paragraph \u2192 type "make this more concise" in the chat')
b('Forgot flood risk details \u2192 type "add the flood risk" and click "Insert at cursor"')
b('Paragraph sounds too informal \u2192 select it, type "formal RICS register"')
b('Need to mention the conservation area \u2192 "add a note about the conservation area status"')

h('Step 5: Final review', 2)
p('The AI assists with drafting \u2014 the professional opinion is yours. Always:')
b('Review every AI-generated paragraph before accepting')
b('Verify facts against your inspection notes')
b('Ensure the value opinion and comparable analysis reflect your professional judgment')
b('Check that no AI-generated text contradicts your observations')

# ============================================================
# IMPORTANT NOTES
# ============================================================
h('Important Notes', 1)

h('The AI is a drafter, not a co-signer', 2)
p('Your name and RICS credentials go on the report. The AI helps you write faster, but every word is your responsibility. The AI will never:')
b('Provide a valuation figure')
b('Override your professional judgment')
b('Insert text without your explicit acceptance (Tab, Insert, or Replace)')

h('Your feedback trains the system', 2)
p('Every time you:')
b('Accept', ' a suggestion (Tab) \u2014 the system learns that was a good suggestion')
b('Dismiss', ' a suggestion (Esc) \u2014 the system learns to avoid that type of suggestion')
b('Accept then edit', ' \u2014 the system learns what you changed and why')
p('Over time, suggestions become more aligned with your writing style and preferences.')

# ============================================================
# TROUBLESHOOTING
# ============================================================
h('Troubleshooting', 1)
table(
    ['Problem', 'Solution'],
    [
        ['No ghost text appears', 'Check "Assist ON" button is green. Check you\'ve typed at least 10 characters.'],
        ['Ghost text is irrelevant', 'Make sure a property is loaded (search and select an address first).'],
        ['Suggestions are slow (>3s)', 'Free-tier API may be rate-limited. Try Ctrl+Space to trigger manually.'],
        ['Tab inserts a tab character', 'This happens when no ghost text is showing. Tab only accepts visible suggestions.'],
        ['Chat returns an error', 'The AI provider may be temporarily unavailable. Wait and try again.'],
        ['Chat doesn\'t know my property', 'Search and select a property before opening the editor.'],
        ['"Replace selection" missing', 'You need text selected (highlighted) in the editor first.'],
    ]
)

# ============================================================
# QUICK REFERENCE
# ============================================================
h('Quick Reference Card', 1)

mono("""GHOST TEXT                          AI CHAT
\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500               \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
Tab         = Accept all            Enter       = Send message
Ctrl+\u2192      = Accept word           Shift+Enter = New line
Esc         = Dismiss
Ctrl+Space  = Request now           "Insert at cursor"  = Add text
Any key     = Dismiss & type        "Replace selection" = Swap text
                                    "Copy"              = Clipboard

TOOLBAR
\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
[\u2728 Assist ON/OFF]    = Toggle ghost text
[\u2697 AI]               = Toggle sidebar""")

doc.add_paragraph()
pa = doc.add_paragraph()
pa.add_run('Vibe Valuation is part of PropVal by AllRange.').italic = True

output = r'C:\Users\licww\Desktop\propval-mvp\Vibe Valuation\Vibe Valuation User Guide.docx'
doc.save(output)
print(f'Saved to {output}')
